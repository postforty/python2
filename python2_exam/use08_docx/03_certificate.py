# 구글에서 "상장 금테 고화질" 검색
# https://mintp.tistory.com/entry/상장-양식
# pip install python-docx
import docx # Word 문서 조작을 위한 주요 라이브러리
from docx.oxml.ns import qn # XML 네임스페이스 관련 모듈
from docx.enum.text import WD_ALIGN_PARAGRAPH # 텍스트 정렬을 위한 열거형

doc = docx.Document("certificate_form.docx")

style = doc.styles["Normal"]
# style.font.name = '맑은 고딕'
# style._element.rPr.rFonts.set(qn('w:eastAsia'), '맑은 고딕')
style.font.size = docx.shared.Pt(12)

para = doc.add_paragraph()
run = para.add_run(" 제 2025-001 호\n")

# 폰트 설정
# run.font.name = "HY궁서"
# run._element.rPr.rFonts.set(qn("w:eastAsia"), "HY궁서") # 한글과 같은 동아시아 언어 폰트를 설정
# run.font.size = docx.shared.Pt(20)

def set_font(run, font_name, font_size, bold=False):
    """텍스트 폰트, 크기, 굵기를 설정하는 함수이다."""
    run.font.name = font_name
    run._element.rPr.rFonts.set(qn("w:eastAsia"), font_name)
    run.font.size = docx.shared.Pt(font_size)
    run.bold = bold

set_font(run, "궁서", 20)

para = doc.add_paragraph()
run = para.add_run("\n")
run = para.add_run("수 료 증")
run.bold = True
run.font.name = "궁서"
run._element.rPr.rFonts.set(qn("w:eastAsia"), "궁서")
run.font.size = docx.shared.Pt(40)
para.alignment = WD_ALIGN_PARAGRAPH.CENTER

para = doc.add_paragraph()
run = para.add_run("\n\n")
run = para.add_run(" 성\u3000\u3000명: 김일남\n")
run.font.name = "맑은 고딕"
run._element.rPr.rFonts.set(qn("w:eastAsia"), "맑은 고딕")
run.font.size = docx.shared.Pt(20)

run = para.add_run(" 생년월일: 1950.01.01\n")
run.font.name = "나눔고딕"
run._element.rPr.rFonts.set(qn("w:eastAsia"), "나눔고딕")
run.font.size = docx.shared.Pt(20)

run = para.add_run(" 교육과정: Python 1, Python 2\n")
run.font.name = "나눔고딕"
run._element.rPr.rFonts.set(qn("w:eastAsia"), "나눔고딕")
run.font.size = docx.shared.Pt(20)

run = para.add_run(" 교육날짜: 2025.03.09 ~ 2025.05.03\n")
run.font.name = "나눔고딕"
run._element.rPr.rFonts.set(qn("w:eastAsia"), "나눔고딕")
run.font.size = docx.shared.Pt(20)

para = doc.add_paragraph()
run = para.add_run("\n")
run = para.add_run(" 위 사람은 Python 1, Python 2 교육과정을\n")
run.font.name = "나눔고딕"
run._element.rPr.rFonts.set(qn("w:eastAsia"), "나눔고딕")
run.font.size = docx.shared.Pt(20)

run = para.add_run(" 탁월하게 이수하였으므로 이 증서를 수여 합니다.\n")
run.font.name = "휴먼옛체"
run._element.rPr.rFonts.set(qn("w:eastAsia"), "휴먼옛체")
run.font.size = docx.shared.Pt(20)

para = doc.add_paragraph()
run = para.add_run("\n\n")
run = para.add_run("코리아 IT 아카데미")
run.font.name = "나눔고딕"
run.bold = True
run._element.rPr.rFonts.set(qn("w:eastAsia"), "나눔고딕")
run.font.size = docx.shared.Pt(20)
para.alignment = WD_ALIGN_PARAGRAPH.CENTER

doc.save("certificate_result.docx")
