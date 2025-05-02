import docx
from docx.oxml.ns import qn
from docx.enum.text import WD_ALIGN_PARAGRAPH

def set_font(run, font_name, font_size, bold=False):
    """텍스트 폰트, 크기, 굵기를 설정하는 함수이다."""
    run.font.name = font_name
    run._element.rPr.rFonts.set(qn("w:eastAsia"), font_name)
    run.font.size = docx.shared.Pt(font_size)
    run.bold = bold

doc = docx.Document("certificate_form.docx")

style = doc.styles["Normal"]
style.font.size = docx.shared.Pt(12)

# 변수 정의
certificate_number = "제 2025-001 호"
certificate_title = "수 료 증"
name = "김일남"
birth_date = "1950.01.01"
courses = "Python 1, Python 2"
education_date = "2025.03.09 ~ 2025.05.03"
award_statement_1 = "위 사람은"
award_statement_2 = "교육과정을 탁월하게 이수하였으므로 이 증서를 수여 합니다."
academy_name = "코리아 IT 아카데미"

para = doc.add_paragraph()
set_font(para.add_run(f"{certificate_number}\n"), "HY궁서", 20)

para = doc.add_paragraph()
set_font(para.add_run(certificate_title), "나눔고딕", 40, bold=True)
para.alignment = WD_ALIGN_PARAGRAPH.CENTER

para = doc.add_paragraph()
set_font(para.add_run(f"\n성\u3000\u3000명: {name}\n생년월일: {birth_date}\n교육과정: {courses}\n교육날짜: {education_date}\n"), "맑은 고딕", 20)

para = doc.add_paragraph()
award_text = f"{award_statement_1} {courses} {award_statement_2}\n"
run1 = para.add_run(award_text)
set_font(run1, "나눔고딕", 20)  # 첫 번째 Run에 "나눔고딕" 스타일 적용

if len(para.runs) > 1:
    para.runs[1].font.name = "휴먼옛체"
    para.runs[1]._element.rPr.rFonts.set(qn("w:eastAsia"), "휴먼옛체")

para = doc.add_paragraph()
set_font(para.add_run(academy_name), "나눔고딕", 20, bold=True)
para.alignment = WD_ALIGN_PARAGRAPH.CENTER

doc.save("certificate_result.docx")
