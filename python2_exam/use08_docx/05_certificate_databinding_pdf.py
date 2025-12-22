# pip install python-docx
import docx
from docx.oxml.ns import qn
from docx.enum.text import WD_ALIGN_PARAGRAPH

# pdf 변환
# pip install docx2pdf
from docx2pdf import convert

from openpyxl import load_workbook

load_wb = load_workbook('certificate.xlsx') # 엑셀 파일 읽어오기
load_ws = load_wb.active # 읽어온 엑셀 파일에서 활성화된 시트 불러오기

name_list = []
birth_date_list = []
number_list = []

for i in range(1, load_ws.max_row + 1):
    name_list.append(load_ws.cell(i, 1).value)
    birth_date_list.append(load_ws.cell(i, 2).value)
    number_list.append(load_ws.cell(i, 3).value)

print(name_list)
print(birth_date_list)
print(number_list)

def set_font(run, font_name, font_size, bold=False):
    """텍스트 폰트, 크기, 굵기를 설정하는 함수이다."""
    run.font.name = font_name
    run._element.rPr.rFonts.set(qn("w:eastAsia"), font_name)
    run.font.size = docx.shared.Pt(font_size)
    run.bold = bold

for i in range(len(name_list)):
    doc = docx.Document("certificate_form.docx")

    style = doc.styles["Normal"]
    style.font.size = docx.shared.Pt(12)

    # 변수 정의
    certificate_number = number_list[i]
    certificate_title = "수 료 증"
    name = name_list[i]
    birth_date = birth_date_list[i]
    courses = "Python 1, Python 2"
    education_date = "2025.03.09 ~ 2025.05.03"
    award_statement_1 = "위 사람은"
    award_statement_2 = "교육과정을 탁월하게 이수하였으므로 이 증서를 수여 합니다."
    academy_name = "코리아 IT 아카데미"

    para = doc.add_paragraph()
    set_font(para.add_run(f"{certificate_number}\n"), "HY궁서", 20)

    para = doc.add_paragraph()
    set_font(para.add_run(certificate_title), "나눔고딕", 40, bold=True)
    para.alignment = WD_ALIGN_PARAGRAPH.CENTER

    para = doc.add_paragraph()
    set_font(para.add_run(f"\n성\u3000\u3000명: {name}\n생년월일: {birth_date}\n교육과정: {courses}\n교육날짜: {education_date}\n"), "맑은 고딕", 20)

    para = doc.add_paragraph()
    award_text = f"{award_statement_1} {courses} {award_statement_2}\n"
    run1 = para.add_run(award_text)
    set_font(run1, "나눔고딕", 20)  # 첫 번째 Run에 "나눔고딕" 스타일 적용

    if len(para.runs) > 1:
        para.runs[1].font.name = "휴먼옛체"
        para.runs[1]._element.rPr.rFonts.set(qn("w:eastAsia"), "휴먼옛체")

    para = doc.add_paragraph()
    set_font(para.add_run(academy_name), "나눔고딕", 20, bold=True)
    para.alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.save(f'certificate_result_{name_list[i]}.docx')
    convert(f'certificate_result_{name_list[i]}.docx', f'certificate_result_{name_list[i]}.pdf')