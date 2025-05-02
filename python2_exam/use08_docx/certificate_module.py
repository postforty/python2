# pip install python-docx
import docx
from docx.oxml.ns import qn
from docx.enum.text import WD_ALIGN_PARAGRAPH

from openpyxl import load_workbook

# pip install docx2pdf
from docx2pdf import convert

# 변수 정의
certificate_title = "Title"
courses = "Coursework"
education_date = "2020.01.01 ~ 2020.02.01"
award_statement_1 = "Hi, there!"
award_statement_2 = "^_____^"
academy_name = "ABCD"

def make_certificate(name, birth_date, certificate_number, form_file_name, output_file_name, **args):

    contents = {}

    if args:
        contents.update(args)

    for key, value in contents.items():
        globals()[key] = value

    def set_font(run, font_name, font_size, bold=False):
        """텍스트 폰트, 크기, 굵기를 설정하는 함수이다."""
        run.font.name = font_name
        run._element.rPr.rFonts.set(qn("w:eastAsia"), font_name)
        run.font.size = docx.shared.Pt(font_size)
        run.bold = bold

    doc = docx.Document(f"{form_file_name}.docx")

    style = doc.styles["Normal"]
    style.font.size = docx.shared.Pt(12)

    para = doc.add_paragraph()
    set_font(para.add_run(f"{certificate_number}\n"), "HY궁서", 20)

    para = doc.add_paragraph()
    set_font(para.add_run(certificate_title), "나눔고딕", 40, bold=True)
    para.alignment = WD_ALIGN_PARAGRAPH.CENTER

    para = doc.add_paragraph()
    set_font(para.add_run(f"\n성\u3000\u3000명: {name}\n생년월일: {birth_date}\n교육과정: {courses}\n교육날짜: {education_date}\n"), "맑은 고딕", 20)

    para = doc.add_paragraph()
    award_text = f"{award_statement_1} {courses} {award_statement_2}\n"
    run1 = para.add_run(award_text)
    set_font(run1, "나눔고딕", 20)  # 첫 번째 Run에 "나눔고딕" 스타일 적용

    if len(para.runs) > 1:
        para.runs[1].font.name = "휴먼옛체"
        para.runs[1]._element.rPr.rFonts.set(qn("w:eastAsia"), "휴먼옛체")

    para = doc.add_paragraph()
    set_font(para.add_run(academy_name), "나눔고딕", 20, bold=True)
    para.alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.save(f'{output_file_name}.docx')
    convert(f'{output_file_name}.docx', f'{output_file_name}.pdf')

if __name__ == "__main__":
    make_certificate("김일남", "1951-01-01", "2025-001", "certificate_form", "certificate_result")