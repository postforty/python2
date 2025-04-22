# pip install python-docx
import docx
from docx.oxml.ns import qn
from docx.enum.text import WD_ALIGN_PARAGRAPH

from openpyxl import load_workbook

# pip install docx2pdf
from docx2pdf import convert

def make_certificate(name_list, birth_date_list, number_list):
    doc = docx.Document('certificate_form.docx')

    style = doc.styles['Normal']
    style.font.name = '나눔고딕'
    style._element.rPr.rFonts.set(qn('w:eastAsia'), '나눔고딕')
    style.font.size = docx.shared.Pt(12)

    para = doc.add_paragraph()
    # run = para.add_run('\n\n')
    run = para.add_run(f' 제 {number_list} 호\n')
    run.font.name = '나눔고딕'
    run._element.rPr.rFonts.set(qn('w:eastAsia'), '나눔고딕')
    run.font.size = docx.shared.Pt(20)

    para = doc.add_paragraph()
    run = para.add_run('\n\n')
    run = para.add_run('수 료 증')
    run.font.name = '나눔고딕'
    run.bold = True
    run._element.rPr.rFonts.set(qn('w:eastAsia'), '나눔고딕')
    run.font.size = docx.shared.Pt(40)
    para.alignment = WD_ALIGN_PARAGRAPH.CENTER

    para = doc.add_paragraph()
    run = para.add_run('\n\n')
    run = para.add_run(f' 성\u3000\u3000명: {name_list}\n')
    run.font.name = '나눔고딕'
    run._element.rPr.rFonts.set(qn('w:eastAsia'), '나눔고딕')
    run.font.size = docx.shared.Pt(20)

    run = para.add_run(f' 생년월일: {birth_date_list}\n')
    run.font.name = '나눔고딕'
    run._element.rPr.rFonts.set(qn('w:eastAsia'), '나눔고딕')
    run.font.size = docx.shared.Pt(20)

    run = para.add_run(' 교육과정: Python 1, Python 2\n')
    run.font.name = '나눔고딕'
    run._element.rPr.rFonts.set(qn('w:eastAsia'), '나눔고딕')
    run.font.size = docx.shared.Pt(20)

    run = para.add_run(' 교육날짜: 2022.12.17 ~ 2023.2.26\n')
    run.font.name = '나눔고딕'
    run._element.rPr.rFonts.set(qn('w:eastAsia'), '나눔고딕')
    run.font.size = docx.shared.Pt(20)

    para = doc.add_paragraph()
    run = para.add_run('\n\n')
    run = para.add_run(' 위 사람은 Python 1, Python 2 교육과정을\n')
    run.font.name = '나눔고딕'
    run._element.rPr.rFonts.set(qn('w:eastAsia'), '나눔고딕')
    run.font.size = docx.shared.Pt(20)

    run = para.add_run(' 탁월하게 이수하였으므로 이 증서를 수여 합니다.\n')
    run.font.name = '나눔고딕'
    run._element.rPr.rFonts.set(qn('w:eastAsia'), '나눔고딕')
    run.font.size = docx.shared.Pt(20)

    para = doc.add_paragraph()
    run = para.add_run('\n\n\n')
    run = para.add_run('코리아 IT 아카데미')
    run.font.name = '나눔고딕'
    run.bold = True
    run._element.rPr.rFonts.set(qn('w:eastAsia'), '나눔고딕')
    run.font.size = docx.shared.Pt(20)
    para.alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.save(f'certificate_result_{name_list}.docx')
    convert(f'certificate_result_{name_list}.docx', f'certificate_result_{name_list}.pdf')

if __name__ == "__main__":
    load_wb = load_workbook('certificate.xlsx') # 엑셀 파일 읽어오기
    load_ws = load_wb.active # 읽어온 엑셀 파일에서 활성화된 시트 불러오기

    name_list = []
    birth_date_list = []
    number_list = []

    for i in range(1, load_ws.max_row + 1):
        name_list.append(load_ws.cell(i, 1).value)
        birth_date_list.append(load_ws.cell(i, 2).value)
        number_list.append(load_ws.cell(i, 3).value)

    print(name_list)
    print(birth_date_list)
    print(number_list)

    for i in range(len(name_list)):
        make_certificate(name_list[i], birth_date_list[i], number_list[i])