# 구글에서 "상장 금테 고화질" 검색
# https://mintp.tistory.com/entry/상장-양식
# pip install python-docx
import docx
from docx.oxml.ns import qn
from docx.enum.text import WD_ALIGN_PARAGRAPH

doc = docx.Document("certificate_form.docx")

style = doc.styles["Normal"]
# style.font.name = '맑은 고딕'
# style._element.rPr.rFonts.set(qn('w:eastAsia'), '맑은 고딕')
style.font.size = docx.shared.Pt(12)

para = doc.add_paragraph()
# run = para.add_run('\n\n')
run = para.add_run(" 제 2020-001 호\n")
run.font.name = "HY궁서"
run._element.rPr.rFonts.set(qn("w:eastAsia"), "HY궁서")
run.font.size = docx.shared.Pt(20)

para = doc.add_paragraph()
run = para.add_run("\n\n")
run = para.add_run("수 료 증")
# run.font.name = '나눔고딕'
run.bold = True
# run._element.rPr.rFonts.set(qn('w:eastAsia'), '나눔고딕')
run.font.size = docx.shared.Pt(40)
para.alignment = WD_ALIGN_PARAGRAPH.CENTER

para = doc.add_paragraph()
run = para.add_run("\n\n")
run = para.add_run(" 성\u3000\u3000명: 김일남\n")
run.font.name = "맑은 고딕"
run._element.rPr.rFonts.set(qn("w:eastAsia"), "맑은 고딕")
run.font.size = docx.shared.Pt(20)

run = para.add_run(" 생년월일: 1950.01.01\n")
run.font.name = "나눔고딕"
run._element.rPr.rFonts.set(qn("w:eastAsia"), "나눔고딕")
run.font.size = docx.shared.Pt(20)

run = para.add_run(" 교육과정: Python 1, Python 2\n")
run.font.name = "나눔고딕"
run._element.rPr.rFonts.set(qn("w:eastAsia"), "나눔고딕")
run.font.size = docx.shared.Pt(20)

run = para.add_run(" 교육날짜: 2022.12.17 ~ 2023.2.26\n")
run.font.name = "나눔고딕"
run._element.rPr.rFonts.set(qn("w:eastAsia"), "나눔고딕")
run.font.size = docx.shared.Pt(20)

para = doc.add_paragraph()
run = para.add_run("\n\n")
run = para.add_run(" 위 사람은 Python 1, Python 2 교육과정을\n")
run.font.name = "나눔고딕"
run._element.rPr.rFonts.set(qn("w:eastAsia"), "나눔고딕")
run.font.size = docx.shared.Pt(20)

run = para.add_run(" 탁월하게 이수하였으므로 이 증서를 수여 합니다.\n")
run.font.name = "휴먼옛체"
run._element.rPr.rFonts.set(qn("w:eastAsia"), "휴먼옛체")
run.font.size = docx.shared.Pt(20)

para = doc.add_paragraph()
run = para.add_run("\n\n\n")
run = para.add_run("코리아 IT 아카데미")
run.font.name = "나눔고딕"
run.bold = True
run._element.rPr.rFonts.set(qn("w:eastAsia"), "나눔고딕")
run.font.size = docx.shared.Pt(20)
para.alignment = WD_ALIGN_PARAGRAPH.RIGHT

print(dir(doc))

doc.save("certificate_result.docx")
